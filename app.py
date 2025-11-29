# import os
# import io
# import streamlit as st
# import numpy as np
# import pydicom
# from PIL import Image as PILImage
# from docx import Document
# from docx.shared import Inches
# import fitz  # PyMuPDF
# import google.generativeai as genai
# from pydicom.pixel_data_handlers.util import apply_modality_lut
# from dotenv import load_dotenv

# # --------- API KEY LOADING ----------
# load_dotenv()
# API_KEY = os.getenv("GOOGLE_API_KEY")

# # --------- STREAMLIT PAGE CONFIG ----------
# st.set_page_config(
#     page_title="MedSight AI Pro",
#     page_icon="🏥",
#     layout="wide",
# )

# # Global styling for a more professional look
# st.markdown(
#     """
#     <style>
#     .block-container { 
#         max-width: 1200px; 
#         padding-top: 2.5rem; 
#     }
#     .medsight-hero {
#         padding: 1.2rem 1.4rem;
#         border-radius: 0.85rem;
#         border: 1px solid #262626;
#         background: linear-gradient(120deg, #111827, #020617);
#         margin-bottom: 1.2rem;
#     }
#     .medsight-hero-title {
#         font-size: 1.25rem;
#         font-weight: 600;
#         margin-bottom: 0.15rem;
#     }
#     .medsight-hero-sub {
#         font-size: 0.9rem;
#         color: #9ca3af;
#     }
#     .report-card {
#         padding: 1.3rem 1.4rem;
#         border-radius: 0.9rem;
#         border: 1px solid #262626;
#         background: #020617;
#         margin-top: 1.0rem;
#         margin-bottom: 0.5rem;
#     }
#     .report-label {
#         font-size: 0.78rem;
#         text-transform: uppercase;
#         letter-spacing: 0.12em;
#         color: #9ca3af;
#         margin-bottom: 0.25rem;
#     }
#     .report-title {
#         font-size: 1.1rem;
#         font-weight: 600;
#         margin-bottom: 0.4rem;
#     }
#     </style>
#     """,
#     unsafe_allow_html=True,
# )

# # --------- DICOM HELPERS ----------
# def get_all_dicom_metadata(dicom_file: pydicom.FileDataset) -> str:
#     lines = ["--- Full DICOM Header Metadata ---"]
#     for tag in dicom_file.iterall():
#         if tag.keyword == "PixelData":
#             continue
#         try:
#             lines.append(f"{tag.name} ({tag.tag}): {tag.value}")
#         except Exception:
#             lines.append(f"{tag.name} ({tag.tag}): [Unreadable]")
#     lines.append("---------------------------------")
#     return "\n".join(lines)


# def handle_dicom_file(uploaded_file):
#     """Return (images, metadata_text) for a DICOM file."""
#     try:
#         dicom_bytes = io.BytesIO(uploaded_file.getvalue())
#         dicom_data = pydicom.dcmread(dicom_bytes)
#         full_metadata_text = get_all_dicom_metadata(dicom_data)
#         images = []

#         try:
#             pixel_array_full = apply_modality_lut(dicom_data.pixel_array, dicom_data)
#         except Exception as e:
#             st.error(f"Error decompressing pixel data: {e}")
#             return [], ""

#         if hasattr(dicom_data, "NumberOfFrames") and dicom_data.NumberOfFrames > 1:
#             frame_indices = [0, dicom_data.NumberOfFrames // 2, dicom_data.NumberOfFrames - 1]
#             pixel_arrays = pixel_array_full
#         else:
#             frame_indices = [0]
#             pixel_arrays = [pixel_array_full]

#         for i in frame_indices:
#             pixel_array = pixel_arrays[i].astype(float)
#             if pixel_array.max() > 0:
#                 rescaled_array = (np.maximum(pixel_array, 0) / pixel_array.max()) * 255
#             else:
#                 rescaled_array = np.zeros_like(pixel_array)

#             final_array = np.uint8(rescaled_array)
#             image = PILImage.fromarray(final_array)
#             if image.mode != "RGB":
#                 image = image.convert("RGB")
#             images.append(image)

#         return images, full_metadata_text

#     except Exception as e:
#         st.error(f"Error processing DICOM file: {e}")
#         return [], ""


# # --------- PDF HELPER ----------
# def handle_pdf_file(uploaded_file):
#     """Return (images, text) for a PDF file."""
#     try:
#         pdf_bytes = uploaded_file.getvalue()
#         doc = fitz.open(stream=pdf_bytes, filetype="pdf")
#         text_content = ""
#         images = []

#         for page in doc:
#             text_content += page.get_text() + "\n\n"

#         for page_num in range(len(doc)):
#             imgs = doc.get_page_images(page_num)
#             if imgs:
#                 xref = imgs[0][0]
#                 base_img = doc.extract_image(xref)
#                 img_bytes = base_img["image"]
#                 img = PILImage.open(io.BytesIO(img_bytes))
#                 if img.mode != "RGB":
#                     img = img.convert("RGB")
#                 images.append(img)
#                 break

#         return images, text_content
#     except Exception as e:
#         st.error(f"Error processing PDF: {e}")
#         return [], ""


# # --------- DOCX REPORT ----------
# def create_docx(doctor_text: str, patient_text: str | None, imgs):
#     doc = Document()
#     doc.add_heading("MedSight AI Report", 0)

#     doc.add_heading("Specialist Report (For Clinicians)", level=1)
#     doc.add_paragraph(doctor_text)

#     if patient_text:
#         doc.add_page_break()
#         doc.add_heading("Patient-Friendly Explanation", level=1)
#         doc.add_paragraph(patient_text)

#     if imgs:
#         doc.add_page_break()
#         doc.add_heading("Reference Images", level=1)
#         for i, img in enumerate(imgs):
#             doc.add_paragraph(f"Image {i + 1}")
#             img_io = io.BytesIO()
#             img.save(img_io, format="PNG")
#             img_io.seek(0)
#             doc.add_picture(img_io, width=Inches(5.0))

#     return doc


# def get_docx_bytes(doctor_text, patient_text, imgs):
#     doc = create_docx(doctor_text, patient_text, imgs)
#     buffer = io.BytesIO()
#     doc.save(buffer)
#     buffer.seek(0)
#     return buffer.getvalue()


# # --------- STUDY TYPE DETECTION ----------
# def detect_study_type(filename: str, text: str) -> str:
#     """
#     Heuristically detect study type from filename + extracted text/metadata.
#     """
#     name = (filename or "").lower()
#     t = (text or "").lower()

#     eye_kw = ["vng", "videonystagmography", "nystagmus", "smooth pursuit",
#               "saccade", "saccadic", "vergence", "vhit", "vor", "oculomotor"]
#     if any(k in t or k in name for k in eye_kw):
#         return "Eye movement / VNG or oculomotor assessment"

#     mri_kw = ["mri", "magnetic resonance", "t1-weighted", "t2-weighted", "flair"]
#     ct_kw = [" ct ", "ct-", "computed tomography", "ct scan"]
#     if any(k in t or k in name for k in mri_kw):
#         return "MRI study"
#     if any(k in t or k in name for k in ct_kw):
#         return "CT study"

#     xray_kw = ["x-ray", "x ray", "radiograph", "cxr", "chest ap", "chest pa",
#                "modality (0008, 0060): cr", "modality (0008, 0060): dx"]
#     if any(k in t or k in name for k in xray_kw):
#         region = ""
#         if "chest" in t or "thorax" in t:
#             region = " chest"
#         elif "spine" in t:
#             region = " spine"
#         return f"X-ray{region} study"

#     if "graph" in t or "trial" in t or "trace" in t:
#         return "Functional / graph-based test report"

#     return "General medical imaging / report"


# # --------- TEXT CLEAN-UP ----------
# CLINICAL_LINE = (
#     "These imaging findings must be correlated with the patient's clinical history, "
#     "examination, and other diagnostic tests before a final diagnosis is made."
# )

# IMAGE_QUALITY_BLOCK = """IMAGE QUALITY

# The submitted image is of non-diagnostic quality. It is severely degraded by an artistic digital filter or overlay, which obscures anatomical detail, creates spurious signals, and prevents reliable assessment. The following findings are based on the limited visible information and should be interpreted with extreme caution."""


# def clean_report(text: str) -> str:
#     """Remove unwanted fixed lines/blocks from the model output."""
#     if not text:
#         return text

#     # Remove the clinical correlation sentence if it appears
#     text = text.replace(CLINICAL_LINE, "").strip()

#     # Remove the specific IMAGE QUALITY block if present
#     text = text.replace(IMAGE_QUALITY_BLOCK, "").strip()

#     # Also remove doubled blank lines created by removals
#     while "\n\n\n" in text:
#         text = text.replace("\n\n\n", "\n\n")

#     return text


# # --------- CORE PROMPT ----------
# BASE_PROMPT = """
# You are a board-certified specialist radiologist and neuro-otology imaging expert.
# You are assisting another clinician by interpreting uploaded medical images and any associated medical reports.

# GENERAL BEHAVIOUR
# - Act as a specialist imaging consultant, not a treating doctor.
# - Focus ONLY on describing and interpreting what can reasonably be inferred from the provided images and text.
# - Do NOT give treatment recommendations, management plans, lifestyle advice, or referrals.
# - Do NOT tell the reader or patient what they "should" do (no "go see your doctor" or similar).
# - Do NOT propose medications, surgeries, or specific procedures.
# - Do NOT use legal or insurance terms such as "apportionment".
# - Only comment on image quality if the context clearly indicates a quality issue. Do NOT invent statements about artistic filters, non-diagnostic quality, or severe degradation if this is not explicit in the input. NEVER use the phrase about artistic digital filters obscuring anatomical detail.

# MULTI-STUDY / FOLLOW-UP LOGIC
# - Multiple studies may be uploaded, sometimes from different dates or modalities.
# - Clearly describe:
#   - Key findings for each study type (e.g., X-ray, CT/MRI, VNG).
#   - Whether findings across studies appear similar, improved, worsened, or indeterminate.
# - If timing or comparability is unclear, explicitly say that change over time cannot be reliably determined.

# EYE MOVEMENT / VNG / VOR (ONLY WHEN RELEVANT)
# - If the context indicates VNG or eye-movement testing, interpret using appropriate oculomotor terminology.
# - Do NOT assume that vestibulo-ocular reflex (VOR) was tested on VNG unless this is clearly stated.
# - If VOR is not directly measured, say: "VOR cannot be directly assessed from these VNG findings alone."
# - For smooth pursuit, when data are available, comment separately for Left, Right, Up, and Down directions.
# - Avoid precise phrases such as "1 beat outside normal" unless explicit numeric normative ranges are given; otherwise use qualitative descriptors like "borderline" or "slightly above typical range".

# CLINICIAN-LEVEL REPORT (DOCTOR SECTION)
# - Use professional radiology / neuro-otology language appropriate to the detected study types.
# - Produce a detailed, structured report with headings such as:
#   - Study / Technique
#   - Image Quality (only if truly evident from context)
#   - Findings (organised by organ, level, or subsystem)
#   - Quantitative Metrics
#   - Comparative / Follow-up comments (if multiple studies)
#   - Impression (key points, concise and numbered)
# - In the Findings and Quantitative Metrics:
#   - Extract and restate any numerical values present in the provided reports or context (e.g., angles, gains, velocities, amplitudes, lesion size in mm, degrees of deviation, latency values).
#   - Present these numbers clearly, preferably in bullet points or short lists, so the clinician can quickly see the key measurements.
#   - Do NOT invent numbers that are not present in the input; if no numbers are given, say that quantitative data are not provided rather than creating them.
# - Highlight clinically important patterns that can be derived from the combination of qualitative findings and available numeric values.

# PATIENT-FRIENDLY SUMMARY (PATIENT SECTION)
# - Write for a non-medical reader at about a 12th-grade level.
# - Use very simple words and short paragraphs.
# - Explain WHAT was seen and, in broad terms, HOW it might relate to common symptoms (e.g., pain, dizziness, shortness of breath).
# - Do NOT give advice or instructions. No "you should...", no "talk to your doctor about...", no lifestyle tips, no recommendations.
# - Do NOT mention statistics or probabilities.

# OUTPUT FORMAT (CRITICAL)
# Return your answer in EXACTLY this structure, using the markers below:

# DOCTOR_REPORT_START
# [Markdown-formatted clinician-level report goes here]
# DOCTOR_REPORT_END

# PATIENT_SUMMARY_START
# [Markdown-formatted simple-language explanation for a patient goes here]
# PATIENT_SUMMARY_END

# Do not use these marker words for anything else.
# """

# # --------- HEADER ----------
# st.title("🏥 MedSight AI Pro")

# st.markdown(
#     """
#     <div class="medsight-hero">
#         <div class="medsight-hero-title">Specialist Imaging Assistant for Clinicians</div>
#         <div class="medsight-hero-sub">
#             Interprets medical imaging and related reports as a specialist radiologist. 
#             Outputs one clinician-level report and an optional patient-friendly summary. 
#             Not a replacement for formal radiology reporting or clinical judgement.
#         </div>
#     </div>
#     """,
#     unsafe_allow_html=True,
# )

# if not API_KEY:
#     st.error("🚨 GOOGLE_API_KEY not found in environment.")
#     st.stop()

# genai.configure(api_key=API_KEY)

# # --------- SIDEBAR ----------
# st.sidebar.title("Workspace")

# uploaded_files = st.sidebar.file_uploader(
#     "Upload medical files (DICOM, PDF, images)",
#     type=["dcm", "dicom", "pdf", "jpg", "jpeg", "png"],
#     accept_multiple_files=True,
#     help="Multiple studies allowed. MedSight will auto-detect study type.",
# )

# analyze_button = st.sidebar.button("Run MedSight Analysis", use_container_width=True, type="primary")

# # --------- SESSION STATE ----------
# if "doctor_report" not in st.session_state:
#     st.session_state["doctor_report"] = ""
# if "patient_summary" not in st.session_state:
#     st.session_state["patient_summary"] = ""
# if "all_images" not in st.session_state:
#     st.session_state["all_images"] = []
# if "preview_data" not in st.session_state:
#     st.session_state["preview_data"] = []

# context_snippets = []
# all_images = []
# detected_types_for_prompt = []
# preview_data = []

# # --------- PROCESS FILES (NO UI YET) ----------
# if uploaded_files:
#     for idx, uploaded_file in enumerate(uploaded_files, start=1):
#         ext = uploaded_file.name.lower().split(".")[-1]
#         label = f"Study {idx}: {uploaded_file.name}"

#         imgs = []
#         extra_text = ""

#         if ext in ["dcm", "dicom"]:
#             imgs, extra_text = handle_dicom_file(uploaded_file)
#         elif ext == "pdf":
#             imgs, extra_text = handle_pdf_file(uploaded_file)
#         else:
#             try:
#                 img = PILImage.open(uploaded_file).convert("RGB")
#                 imgs = [img]
#                 extra_text = ""
#             except Exception as e:
#                 st.warning(f"Could not open {uploaded_file.name} as image: {e}")

#         if extra_text:
#             snippet_body = extra_text[:2000]
#         else:
#             snippet_body = "(No additional text extracted; interpretation must rely mainly on imaging appearance.)"

#         snippet_header = f"### {label}\n"
#         context_snippets.append(snippet_header + snippet_body)
#         all_images.extend(imgs[:3])  # cap per file

#         study_type_label = detect_study_type(uploaded_file.name, extra_text)
#         detected_types_for_prompt.append(f"{label}: {study_type_label}")

#         preview_data.append({"label": label, "imgs": imgs})

#     st.session_state["all_images"] = all_images
#     st.session_state["preview_data"] = preview_data

# # --------- RUN ANALYSIS ----------
# if analyze_button:
#     if not uploaded_files:
#         st.warning("Please upload at least one file before analyzing.")
#     else:
#         with st.spinner("Analyzing with Gemini (gemini-pro-latest)..."):
#             try:
#                 # Use gemini-pro-latest as requested
#                 model = genai.GenerativeModel(model_name="gemini-pro-latest")

#                 combined_context = "\n\n".join(context_snippets)
#                 if len(combined_context) > 8000:
#                     combined_context = combined_context[:8000] + "\n\n...[context trimmed for length]"

#                 resized_images = []
#                 for img in all_images:
#                     img_copy = img.copy()
#                     img_copy.thumbnail((512, 512))
#                     resized_images.append(img_copy)

#                 if detected_types_for_prompt:
#                     types_text = "Detected study types:\n" + "\n".join(
#                         f"- {s}" for s in detected_types_for_prompt
#                     )
#                 else:
#                     types_text = "Detected study types: not clearly identifiable; treat as general medical imaging / reports."

#                 preamble = BASE_PROMPT + "\n\n" + types_text + "\n"

#                 inputs = [
#                     preamble,
#                     "--- START OF CONTEXT ---",
#                     combined_context if combined_context.strip() else "No textual context; rely on visual appearance.",
#                     "--- END OF CONTEXT ---",
#                 ]
#                 inputs.extend(resized_images)

#                 token_info = model.count_tokens(inputs)
#                 if token_info.total_tokens > 900000:
#                     st.error("❌ Input too large! Please upload fewer or smaller files.")
#                     st.stop()

#                 response = model.generate_content(inputs)
#                 raw_text = response.text.strip()

#                 doc_start = "DOCTOR_REPORT_START"
#                 doc_end = "DOCTOR_REPORT_END"
#                 pat_start = "PATIENT_SUMMARY_START"
#                 pat_end = "PATIENT_SUMMARY_END"

#                 doctor_report = ""
#                 patient_summary = ""

#                 if doc_start in raw_text and doc_end in raw_text:
#                     doctor_report = raw_text.split(doc_start, 1)[1].split(doc_end, 1)[0].strip()
#                 else:
#                     doctor_report = raw_text

#                 if pat_start in raw_text and pat_end in raw_text:
#                     patient_summary = raw_text.split(pat_start, 1)[1].split(pat_end, 1)[0].strip()
#                 else:
#                     patient_summary = ""

#                 # Clean unwanted lines/blocks
#                 doctor_report = clean_report(doctor_report)
#                 patient_summary = clean_report(patient_summary)

#                 st.session_state["doctor_report"] = doctor_report
#                 st.session_state["patient_summary"] = patient_summary

#             except Exception as e:
#                 st.error(f"An error occurred during analysis: {e}")

# # --------- DISPLAY REPORTS (TOP) ----------
# if st.session_state.get("doctor_report"):
#     st.markdown(
#         """
#         <div class="report-card">
#             <div class="report-label">Clinician Section</div>
#             <div class="report-title">Specialist Doctor Report</div>
#         </div>
#         """,
#         unsafe_allow_html=True,
#     )
#     st.markdown(st.session_state["doctor_report"])

#     show_patient = st.toggle("Show Patient-Friendly Summary", value=False)
#     if show_patient:
#         st.markdown(
#             """
#             <div class="report-card">
#                 <div class="report-label">Patient Section</div>
#                 <div class="report-title">Patient-Friendly Summary</div>
#             </div>
#             """,
#             unsafe_allow_html=True,
#         )
#         if st.session_state.get("patient_summary"):
#             st.markdown(st.session_state["patient_summary"])
#         else:
#             st.info("No separate patient summary was generated for this study.")

#     docx_bytes = get_docx_bytes(
#         st.session_state["doctor_report"],
#         st.session_state.get("patient_summary") or None,
#         imgs=st.session_state.get("all_images"),
#     )

#     st.download_button(
#         "⬇️ Download Full Report (DOCX)",
#         data=docx_bytes,
#         file_name="MedSight_Report.docx",
#         mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
#     )

# # --------- UPLOADED FILES PREVIEW (BOTTOM) ----------
# preview_data = st.session_state.get("preview_data", [])
# if preview_data:
#     with st.expander("📁 Uploaded Files & Preview", expanded=False):
#         for item in preview_data:
#             label = item["label"]
#             imgs = item["imgs"]
#             st.markdown(f"**{label}**")
#             if imgs:
#                 st.image(
#                     imgs,
#                     caption=[f"{label} - Image {i+1}" for i in range(len(imgs))],
#                     width=220,
#                 )


import os
import io
import streamlit as st
import numpy as np
import pydicom
from PIL import Image as PILImage
from docx import Document
from docx.shared import Inches
import fitz  # PyMuPDF
import google.generativeai as genai
from pydicom.pixel_data_handlers.util import apply_modality_lut
from dotenv import load_dotenv

# --------- API KEY LOADING ----------
load_dotenv()
API_KEY = os.getenv("GOOGLE_API_KEY")

# --------- STREAMLIT PAGE CONFIG ----------
st.set_page_config(
    page_title="MedSight",
    page_icon="🏥",
    layout="wide",
)

# --------- GLOBAL STYLING ----------
st.markdown(
    """
    <style>
    :root {
        --med-bg: #020617;
        --med-bg-soft: #050816;
        --med-card: #050816;
        --med-card-soft: #0b1220;
        --med-border: #1f2937;
        --med-accent: #10b981;   /* teal/green */
        --med-accent-soft: rgba(16,185,129,0.12);
        --med-text-main: #e5e7eb;
        --med-text-muted: #9ca3af;
        --med-danger: #ef4444;
    }

    body {
        background: radial-gradient(circle at top left, #111827 0, #020617 45%, #020617 100%);
        color: var(--med-text-main);
    }

    /* Main container */
    .block-container { 
        max-width: 1180px; 
        padding-top: 1.8rem; 
        padding-bottom: 2.5rem;
    }

    /* App title */
    h1 {
        font-weight: 700 !important;
        letter-spacing: 0.04em;
        margin-bottom: 0.2rem;
    }

    /* Cards */
    .workspace-card,
    .report-card {
        padding: 1.2rem 1.4rem;
        border-radius: 1rem;
        border: 1px solid var(--med-border);
        background: radial-gradient(circle at top left, var(--med-card-soft), var(--med-card));
        box-shadow: 0 18px 40px rgba(0,0,0,0.55);
    }

    .workspace-card {
        margin-bottom: 1.0rem;
    }

    .report-card {
        margin-top: 1.1rem;
        margin-bottom: 0.6rem;
    }

    .workspace-title {
        font-size: 1.1rem;
        font-weight: 600;
        margin-bottom: 0.25rem;
    }

    .workspace-sub {
        font-size: 0.88rem;
        color: var(--med-text-muted);
        margin-bottom: 0.9rem;
    }

    .report-label {
        font-size: 0.76rem;
        text-transform: uppercase;
        letter-spacing: 0.16em;
        color: var(--med-accent);
        margin-bottom: 0.25rem;
    }

    .report-title {
        font-size: 1.1rem;
        font-weight: 600;
        margin-bottom: 0.35rem;
    }

    /* File uploader */
    .stFileUploader {
        background: var(--med-bg-soft);
        border-radius: 0.9rem;
        border: 1px dashed var(--med-border);
        padding: 0.6rem 0.9rem;
    }

    /* Buttons */
    div.stButton > button {
        background: linear-gradient(135deg, var(--med-accent), #22c55e);
        color: #0b1120;
        border-radius: 999px;
        border: none;
        font-weight: 600;
        padding: 0.6rem 1.2rem;
        box-shadow: 0 10px 25px rgba(16,185,129,0.35);
    }
    div.stButton > button:hover {
        filter: brightness(1.05);
        box-shadow: 0 14px 30px rgba(16,185,129,0.45);
    }

    /* Expander */
    .streamlit-expanderHeader {
        background: var(--med-card-soft);
        color: var(--med-text-main);
        border-radius: 0.7rem;
    }
    .streamlit-expanderContent {
        background: var(--med-bg-soft);
        border-radius: 0 0 0.7rem 0.7rem;
    }

    .footer-text {
        font-size: 0.85rem;
        color: var(--med-text-muted);
    }

    .stMarkdown p {
        line-height: 1.55;
        font-size: 0.95rem;
    }
    </style>
    """,
    unsafe_allow_html=True,
)

# --------- DICOM HELPERS ----------
def get_all_dicom_metadata(dicom_file: pydicom.FileDataset) -> str:
    lines = ["--- Full DICOM Header Metadata ---"]
    for tag in dicom_file.iterall():
        if tag.keyword == "PixelData":
            continue
        try:
            lines.append(f"{tag.name} ({tag.tag}): {tag.value}")
        except Exception:
            lines.append(f"{tag.name} ({tag.tag}): [Unreadable]")
    lines.append("---------------------------------")
    return "\n".join(lines)


def handle_dicom_file(uploaded_file):
    """Return (images, metadata_text) for a DICOM file."""
    try:
        dicom_bytes = io.BytesIO(uploaded_file.getvalue())
        dicom_data = pydicom.dcmread(dicom_bytes)
        full_metadata_text = get_all_dicom_metadata(dicom_data)
        images = []

        try:
            pixel_array_full = apply_modality_lut(dicom_data.pixel_array, dicom_data)
        except Exception as e:
            st.error(f"Error decompressing pixel data: {e}")
            return [], ""

        if hasattr(dicom_data, "NumberOfFrames") and dicom_data.NumberOfFrames > 1:
            frame_indices = [0, dicom_data.NumberOfFrames // 2, dicom_data.NumberOfFrames - 1]
            pixel_arrays = pixel_array_full
        else:
            frame_indices = [0]
            pixel_arrays = [pixel_array_full]

        for i in frame_indices:
            pixel_array = pixel_arrays[i].astype(float)
            if pixel_array.max() > 0:
                rescaled_array = (np.maximum(pixel_array, 0) / pixel_array.max()) * 255
            else:
                rescaled_array = np.zeros_like(pixel_array)

            final_array = np.uint8(rescaled_array)
            image = PILImage.fromarray(final_array)
            if image.mode != "RGB":
                image = image.convert("RGB")
            images.append(image)

        return images, full_metadata_text

    except Exception as e:
        st.error(f"Error processing DICOM file: {e}")
        return [], ""


# --------- PDF HELPER ----------
def handle_pdf_file(uploaded_file):
    """Return (images, text) for a PDF file."""
    try:
        pdf_bytes = uploaded_file.getvalue()
        doc = fitz.open(stream=pdf_bytes, filetype="pdf")
        text_content = ""
        images = []

        for page in doc:
            text_content += page.get_text() + "\n\n"

        for page_num in range(len(doc)):
            imgs = doc.get_page_images(page_num)
            if imgs:
                xref = imgs[0][0]
                base_img = doc.extract_image(xref)
                img_bytes = base_img["image"]
                img = PILImage.open(io.BytesIO(img_bytes))
                if img.mode != "RGB":
                    img = img.convert("RGB")
                images.append(img)
                break

        return images, text_content
    except Exception as e:
        st.error(f"Error processing PDF: {e}")
        return [], ""


# --------- DOCX REPORT ----------
def create_docx(doctor_text: str, patient_text: str | None, imgs):
    doc = Document()
    doc.add_heading("MedSight Report", 0)

    doc.add_heading("Specialist Report (For Clinicians)", level=1)
    doc.add_paragraph(doctor_text)

    if patient_text:
        doc.add_page_break()
        doc.add_heading("Patient-Friendly Explanation", level=1)
        doc.add_paragraph(patient_text)

    if imgs:
        doc.add_page_break()
        doc.add_heading("Reference Images", level=1)
        for i, img in enumerate(imgs):
            doc.add_paragraph(f"Image {i + 1}")
            img_io = io.BytesIO()
            img.save(img_io, format="PNG")
            img_io.seek(0)
            doc.add_picture(img_io, width=Inches(5.0))

    return doc


def get_docx_bytes(doctor_text, patient_text, imgs):
    doc = create_docx(doctor_text, patient_text, imgs)
    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.getvalue()


# --------- STUDY TYPE DETECTION ----------
def detect_study_type(filename: str, text: str) -> str:
    """
    Heuristically detect study type from filename + extracted text/metadata.
    """
    name = (filename or "").lower()
    t = (text or "").lower()

    eye_kw = ["vng", "videonystagmography", "nystagmus", "smooth pursuit",
              "saccade", "saccadic", "vergence", "vhit", "vor", "oculomotor"]
    if any(k in t or k in name for k in eye_kw):
        return "Eye movement / VNG or oculomotor assessment"

    mri_kw = ["mri", "magnetic resonance", "t1-weighted", "t2-weighted", "flair"]
    ct_kw = [" ct ", "ct-", "computed tomography", "ct scan"]
    if any(k in t or k in name for k in mri_kw):
        return "MRI study"
    if any(k in t or k in name for k in ct_kw):
        return "CT study"

    xray_kw = ["x-ray", "x ray", "radiograph", "cxr", "chest ap", "chest pa",
               "modality (0008, 0060): cr", "modality (0008, 0060): dx"]
    if any(k in t or k in name for k in xray_kw):
        region = ""
        if "chest" in t or "thorax" in t:
            region = " chest"
        elif "spine" in t:
            region = " spine"
        return f"X-ray{region} study"

    if "graph" in t or "trial" in t or "trace" in t:
        return "Functional / graph-based test report"

    return "General medical imaging / report"


# --------- TEXT CLEAN-UP ----------
CLINICAL_LINE = (
    "These imaging findings must be correlated with the patient's clinical history, "
    "examination, and other diagnostic tests before a final diagnosis is made."
)

IMAGE_QUALITY_BLOCK = """IMAGE QUALITY

The submitted image is of non-diagnostic quality. It is severely degraded by an artistic digital filter or overlay, which obscures anatomical detail, creates spurious signals, and prevents reliable assessment. The following findings are based on the limited visible information and should be interpreted with extreme caution."""


def clean_report(text: str) -> str:
    """Remove unwanted fixed lines/blocks from the model output."""
    if not text:
        return text

    text = text.replace(CLINICAL_LINE, "").strip()
    text = text.replace(IMAGE_QUALITY_BLOCK, "").strip()

    while "\n\n\n" in text:
        text = text.replace("\n\n\n", "\n\n")

    return text


# --------- CORE PROMPT ----------
BASE_PROMPT = """
You are a board-certified specialist radiologist and neuro-otology imaging expert.
You are assisting another clinician by interpreting uploaded medical images and any associated medical reports.

GENERAL BEHAVIOUR
- Act as a specialist imaging consultant, not a treating doctor.
- Focus ONLY on describing and interpreting what can reasonably be inferred from the provided images and text.
- Do NOT give treatment recommendations, management plans, lifestyle advice, or referrals.
- Do NOT tell the reader or patient what they "should" do (no "go see your doctor" or similar).
- Do NOT propose medications, surgeries, or specific procedures.
- Do NOT use legal or insurance terms such as "apportionment".
- Only comment on image quality if the context clearly indicates a quality issue. Do NOT invent statements about artistic filters, non-diagnostic quality, or severe degradation if this is not explicit in the input. NEVER use the phrase about artistic digital filters obscuring anatomical detail.

MULTI-STUDY / FOLLOW-UP LOGIC
- Multiple studies may be uploaded, sometimes from different dates or modalities.
- Clearly describe:
  - Key findings for each study type (e.g., X-ray, CT/MRI, VNG).
  - Whether findings across studies appear similar, improved, worsened, or indeterminate.
- If timing or comparability is unclear, explicitly say that change over time cannot be reliably determined.

EYE MOVEMENT / VNG / VOR (ONLY WHEN RELEVANT)
- If the context indicates VNG or eye-movement testing, interpret using appropriate oculomotor terminology.
- Do NOT assume that vestibulo-ocular reflex (VOR) was tested on VNG unless this is clearly stated.
- If VOR is not directly measured, say: "VOR cannot be directly assessed from these VNG findings alone."
- For smooth pursuit, when data are available, comment separately for Left, Right, Up, and Down directions.
- Avoid precise phrases such as "1 beat outside normal" unless explicit numeric normative ranges are given; otherwise use qualitative descriptors like "borderline" or "slightly above typical range".

CLINICIAN-LEVEL REPORT (DOCTOR SECTION)
- Use professional radiology / neuro-otology language appropriate to the detected study types.
- Produce a detailed, structured report with headings such as:
  - Study / Technique
  - Image Quality (only if truly evident from context)
  - Findings (organised by organ, level, or subsystem)
  - Quantitative Metrics
  - Comparative / Follow-up comments (if multiple studies)
  - Impression (key points, concise and numbered)
- In the Findings and Quantitative Metrics:
  - Extract and restate any numerical values present in the provided reports or context (e.g., angles, gains, velocities, amplitudes, lesion size in mm, degrees of deviation, latency values).
  - Present these numbers clearly, preferably in bullet points or short lists, so the clinician can quickly see the key measurements.
  - Do NOT invent numbers that are not present in the input; if no numbers are given, say that quantitative data are not provided rather than creating them.
- Highlight clinically important patterns that can be derived from the combination of qualitative findings and available numeric values.

PATIENT-FRIENDLY SUMMARY (PATIENT SECTION)
- Write for a non-medical reader at about a 12th-grade level.
- Use very simple words and short paragraphs.
- Explain WHAT was seen and, in broad terms, HOW it might relate to common symptoms (e.g., pain, dizziness, shortness of breath).
- Do NOT give advice or instructions. No "you should...", no "talk to your doctor about...", no lifestyle tips, no recommendations.
- Do NOT mention statistics or probabilities.

OUTPUT FORMAT (CRITICAL)
Return your answer in EXACTLY this structure, using the markers below:

DOCTOR_REPORT_START
[Markdown-formatted clinician-level report goes here]
DOCTOR_REPORT_END

PATIENT_SUMMARY_START
[Markdown-formatted simple-language explanation for a patient goes here]
PATIENT_SUMMARY_END

Do not use these marker words for anything else.
"""

# --------- HEADER ----------
st.title("🏥 MedSight")

if not API_KEY:
    st.error("🚨 GOOGLE_API_KEY not found in environment.")
    st.stop()

genai.configure(api_key=API_KEY)

# --------- SESSION STATE ----------
if "doctor_report" not in st.session_state:
    st.session_state["doctor_report"] = ""
if "patient_summary" not in st.session_state:
    st.session_state["patient_summary"] = ""
if "all_images" not in st.session_state:
    st.session_state["all_images"] = []
if "preview_data" not in st.session_state:
    st.session_state["preview_data"] = []

context_snippets = []
all_images = []
detected_types_for_prompt = []
preview_data = []

# --------- WORKSPACE CARD (NO SIDEBAR) ----------
with st.container():
    st.markdown(
        """
        <div class="workspace-card">
            <div class="workspace-title">Workspace</div>
            <div class="workspace-sub">
                Upload one or more medical files. MedSight will auto-detect study type and generate a specialist report and an optional patient summary.
            </div>
        </div>
        """,
        unsafe_allow_html=True,
    )

    col1, col2 = st.columns([3, 1])

    with col1:
        uploaded_files = st.file_uploader(
            "Upload medical files (DICOM, PDF, images)",
            type=["dcm", "dicom", "pdf", "jpg", "jpeg", "png"],
            accept_multiple_files=True,
            help="Multiple studies allowed.",
        )

    with col2:
        st.write("")
        st.write("")
        analyze_button = st.button(
            "Run MedSight Analysis",
            use_container_width=True,
            type="primary",
        )

# --------- PROCESS FILES ----------
if uploaded_files:
    for idx, uploaded_file in enumerate(uploaded_files, start=1):
        ext = uploaded_file.name.lower().split(".")[-1]
        label = f"Study {idx}: {uploaded_file.name}"

        imgs = []
        extra_text = ""

        if ext in ["dcm", "dicom"]:
            imgs, extra_text = handle_dicom_file(uploaded_file)
        elif ext == "pdf":
            imgs, extra_text = handle_pdf_file(uploaded_file)
        else:
            try:
                img = PILImage.open(uploaded_file).convert("RGB")
                imgs = [img]
                extra_text = ""
            except Exception as e:
                st.warning(f"Could not open {uploaded_file.name} as image: {e}")

        if extra_text:
            snippet_body = extra_text[:2000]
        else:
            snippet_body = "(No additional text extracted; interpretation must rely mainly on imaging appearance.)"

        snippet_header = f"### {label}\n"
        context_snippets.append(snippet_header + snippet_body)
        all_images.extend(imgs[:3])  # cap per file

        study_type_label = detect_study_type(uploaded_file.name, extra_text)
        detected_types_for_prompt.append(f"{label}: {study_type_label}")

        preview_data.append({"label": label, "imgs": imgs})

    st.session_state["all_images"] = all_images
    st.session_state["preview_data"] = preview_data

# --------- RUN ANALYSIS ----------
if analyze_button:
    if not uploaded_files:
        st.warning("Please upload at least one file before analyzing.")
    else:
        with st.spinner("Analyzing with Gemini (gemini-pro-latest)..."):
            try:
                model = genai.GenerativeModel(model_name="gemini-pro-latest")

                combined_context = "\n\n".join(context_snippets)
                if len(combined_context) > 8000:
                    combined_context = combined_context[:8000] + "\n\n...[context trimmed for length]"

                resized_images = []
                for img in all_images:
                    img_copy = img.copy()
                    img_copy.thumbnail((512, 512))
                    resized_images.append(img_copy)

                if detected_types_for_prompt:
                    types_text = "Detected study types:\n" + "\n".join(
                        f"- {s}" for s in detected_types_for_prompt
                    )
                else:
                    types_text = "Detected study types: not clearly identifiable; treat as general medical imaging / reports."

                preamble = BASE_PROMPT + "\n\n" + types_text + "\n"

                inputs = [
                    preamble,
                    "--- START OF CONTEXT ---",
                    combined_context if combined_context.strip() else "No textual context; rely on visual appearance.",
                    "--- END OF CONTEXT ---",
                ]
                inputs.extend(resized_images)

                token_info = model.count_tokens(inputs)
                if token_info.total_tokens > 900000:
                    st.error("❌ Input too large! Please upload fewer or smaller files.")
                    st.stop()

                response = model.generate_content(inputs)
                raw_text = response.text.strip()

                doc_start = "DOCTOR_REPORT_START"
                doc_end = "DOCTOR_REPORT_END"
                pat_start = "PATIENT_SUMMARY_START"
                pat_end = "PATIENT_SUMMARY_END"

                doctor_report = ""
                patient_summary = ""

                if doc_start in raw_text and doc_end in raw_text:
                    doctor_report = raw_text.split(doc_start, 1)[1].split(doc_end, 1)[0].strip()
                else:
                    doctor_report = raw_text

                if pat_start in raw_text and pat_end in raw_text:
                    patient_summary = raw_text.split(pat_start, 1)[1].split(pat_end, 1)[0].strip()
                else:
                    patient_summary = ""

                doctor_report = clean_report(doctor_report)
                patient_summary = clean_report(patient_summary)

                st.session_state["doctor_report"] = doctor_report
                st.session_state["patient_summary"] = patient_summary

            except Exception as e:
                st.error(f"An error occurred during analysis: {e}")

# --------- DISPLAY REPORTS ----------
if st.session_state.get("doctor_report"):
    st.markdown(
        """
        <div class="report-card">
            <div class="report-label">Clinician Section</div>
            <div class="report-title">Specialist Doctor Report</div>
        </div>
        """,
        unsafe_allow_html=True,
    )
    st.markdown(st.session_state["doctor_report"])

    show_patient = st.toggle("Show Patient-Friendly Summary", value=False)
    if show_patient:
        st.markdown(
            """
            <div class="report-card">
                <div class="report-label">Patient Section</div>
                <div class="report-title">Patient-Friendly Summary</div>
            </div>
            """,
            unsafe_allow_html=True,
        )
        if st.session_state.get("patient_summary"):
            st.markdown(st.session_state["patient_summary"])
        else:
            st.info("No separate patient summary was generated for this study.")

    docx_bytes = get_docx_bytes(
        st.session_state["doctor_report"],
        st.session_state.get("patient_summary") or None,
        imgs=st.session_state.get("all_images"),
    )

    st.download_button(
        "⬇️ Download Full Report (DOCX)",
        data=docx_bytes,
        file_name="MedSight_Report.docx",
        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )

# --------- UPLOADED FILES PREVIEW ----------
preview_data = st.session_state.get("preview_data", [])
if preview_data:
    with st.expander("📁 Uploaded Files & Preview", expanded=False):
        for item in preview_data:
            label = item["label"]
            imgs = item["imgs"]
            st.markdown(f"**{label}**")
            if imgs:
                st.image(
                    imgs,
                    caption=[f"{label} - Image {i+1}" for i in range(len(imgs))],
                    width=220,
                )

# --------- FOOTER ----------
st.markdown("---")
st.markdown(
    """
    <div class="footer-text">
    <strong>Specialist imaging assistant for clinicians.</strong>
    Interprets medical imaging and related reports as a specialist radiologist.
    Outputs one clinician-level report and an optional patient-friendly summary.
    Not a replacement for formal radiology reporting or clinical judgement.
    </div>
    """,
    unsafe_allow_html=True,
)
